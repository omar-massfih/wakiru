"""Turning uploaded files and fetched pages into plain text for ingestion.

The docs pipeline (:mod:`.store`) only ever sees text; this module is the
boundary where PDFs, DOCX files, and HTML become that text. Extraction is
dispatch-by-content: the file's magic bytes win over its name, so a mislabelled
upload still extracts (or fails) honestly.

URL ingestion is opt-in (``enable_docs_url_ingest``): on a server reachable by
others, "fetch any URL I'm told to" is an SSRF primitive, so it defaults off.
"""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser

from .. import netguard

# Pages larger than this are cut off rather than buffered without bound.
_MAX_FETCH_BYTES = 5_000_000
_FETCH_TIMEOUT_SECONDS = 30

# Tags whose content is never prose, and tags that end a block of prose.
_SKIP_TAGS = {"script", "style", "noscript", "template", "head"}
_BLOCK_TAGS = {
    "p", "div", "br", "li", "ul", "ol", "table", "tr", "h1", "h2", "h3",
    "h4", "h5", "h6", "section", "article", "header", "footer", "blockquote",
}


class ExtractionError(ValueError):
    """Raised when a file or page cannot be turned into text."""


def extract_text(filename: str, content: bytes) -> str:
    """Plain text from an uploaded file (PDF, DOCX, or anything text-like)."""
    if content.startswith(b"%PDF") or filename.lower().endswith(".pdf"):
        return _from_pdf(content)
    # DOCX is a zip (PK…) — but so are many formats, so require the extension.
    if filename.lower().endswith(".docx"):
        return _from_docx(content)
    if filename.lower().endswith((".html", ".htm")):
        return html_to_text(content.decode("utf-8", errors="replace"))
    return content.decode("utf-8", errors="replace")


def _from_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except PyPdfError as exc:
        raise ExtractionError(f"could not read PDF: {exc}") from exc
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text:
        raise ExtractionError("PDF contained no extractable text (scanned images?)")
    return text


def _from_docx(content: bytes) -> str:
    import zipfile

    import docx
    from docx.opc.exceptions import OpcError

    try:
        document = docx.Document(io.BytesIO(content))
    except (OpcError, KeyError, ValueError, zipfile.BadZipFile) as exc:
        raise ExtractionError(f"could not read DOCX: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(p for p in parts if p.strip())
    if not text:
        raise ExtractionError("DOCX contained no text")
    return text


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in _SKIP_TAGS:
            self._skip_depth += 1
        elif tag in _BLOCK_TAGS:
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in _SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1
        elif tag in _BLOCK_TAGS:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self._chunks.append(data)

    @property
    def text(self) -> str:
        raw = "".join(self._chunks)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in raw.splitlines()]
        return "\n".join(line for line in lines if line)


def html_to_text(html: str) -> str:
    """Visible prose from an HTML page (stdlib parser; scripts/styles dropped)."""
    parser = _TextExtractor()
    parser.feed(html)
    parser.close()
    return parser.text


def fetch_url_text(url: str) -> tuple[str, str]:
    """Fetch ``url`` and return ``(title, text)``; HTML is reduced to prose.

    The scheme must be http(s) and the host must resolve to a public address
    (every redirect hop re-checked — see :mod:`assistant.netguard`). Callers
    gate this behind ``enable_docs_url_ingest``.
    """
    try:
        with netguard.urlopen_public(
            url,
            timeout=_FETCH_TIMEOUT_SECONDS,
            headers={"User-Agent": "wakiru-assistant"},
        ) as response:
            content_type = response.headers.get_content_type()
            body = response.read(_MAX_FETCH_BYTES)
    except netguard.BlockedURLError as exc:
        raise ExtractionError(str(exc)) from exc
    except OSError as exc:  # URLError subclasses OSError
        raise ExtractionError(f"could not fetch {url}: {exc}") from exc

    if content_type == "application/pdf":
        return url, _from_pdf(body)
    text = body.decode("utf-8", errors="replace")
    if "html" in content_type:
        title_match = re.search(r"<title[^>]*>(.*?)</title>", text, re.I | re.S)
        title = title_match.group(1).strip() if title_match else url
        return title, html_to_text(text)
    return url, text
